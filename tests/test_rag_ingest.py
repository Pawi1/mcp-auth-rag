"""Tests for rag_ingest.py: format sniffing, PDF/DOCX/text parsing (heading
detection), hierarchical chunking (section boundaries + overlap), and the
ingest_document() orchestration (with rag_store/rag_embed mocked)."""

import io
from unittest.mock import AsyncMock

import fitz
import pytest
from docx import Document as DocxDocument

import rag_ingest as ri


def _make_pdf(sections: list[tuple[str, int, str]]) -> bytes:
    """sections: [(heading, n_words, word_prefix), ...] — one page per section."""
    doc = fitz.open()
    for heading, n_words, prefix in sections:
        page = doc.new_page()
        page.insert_text((72, 72), heading, fontsize=18)
        body = " ".join(f"{prefix}{i}" for i in range(n_words))
        rc = page.insert_textbox(fitz.Rect(72, 100, 523, 780), body, fontsize=9)
        assert rc >= 0, f"test fixture text overflowed the page (rc={rc})"
    data = doc.tobytes()
    doc.close()
    return data


class TestSniffFormat:
    def test_recognizes_supported_extensions(self):
        assert ri.sniff_format("paper.pdf") == "pdf"
        assert ri.sniff_format("paper.PDF") == "pdf"
        assert ri.sniff_format("notes.docx") == "docx"
        assert ri.sniff_format("readme.md") == "md"
        assert ri.sniff_format("plain.txt") == "txt"

    def test_rejects_unsupported_extensions(self):
        assert ri.sniff_format("virus.exe") is None
        assert ri.sniff_format("no_extension") is None


class TestContentHash:
    def test_stable_for_same_bytes(self):
        data = b"some pdf bytes"
        assert ri.content_hash(data) == ri.content_hash(data)

    def test_differs_for_different_bytes(self):
        assert ri.content_hash(b"a") != ri.content_hash(b"b")


class TestParsePdf:
    def test_detects_headings_by_font_size(self):
        pdf = _make_pdf([("Wprowadzenie", 200, "slowo"), ("Metodologia badan", 200, "inneslowo")])
        paragraphs, page_count = ri.parse_pdf(pdf)
        assert page_count == 2
        headings = {p.heading for p in paragraphs}
        assert headings == {"Wprowadzenie", "Metodologia badan"}

    def test_paragraphs_carry_their_page_number(self):
        pdf = _make_pdf([("Sekcja A", 50, "a"), ("Sekcja B", 50, "b")])
        paragraphs, _ = ri.parse_pdf(pdf)
        pages_for_a = {p.page for p in paragraphs if p.heading == "Sekcja A"}
        pages_for_b = {p.page for p in paragraphs if p.heading == "Sekcja B"}
        assert pages_for_a == {1}
        assert pages_for_b == {2}

    def test_oversized_symbol_between_paragraphs_does_not_become_a_heading(self):
        # equation operators (∑, ∏...) often render oversized in scientific
        # PDFs — a lone one must not get misdetected as a new section,
        # which would silently mislabel every paragraph after it.
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Wyniki", fontsize=18)
        page.insert_textbox(fitz.Rect(72, 100, 523, 250), "Pierwszy akapit tekstu.", fontsize=9)
        page.insert_text((72, 270), "#", fontsize=30)  # oversized, short, non-alpha — same shape as a stray "∑"
        page.insert_textbox(fitz.Rect(72, 320, 523, 500), "Drugi akapit po symbolu.", fontsize=9)
        pdf_bytes = doc.tobytes()
        doc.close()

        paragraphs, _ = ri.parse_pdf(pdf_bytes)
        assert {p.heading for p in paragraphs} == {"Wyniki"}


class TestLooksLikeHeading:
    def test_lone_math_symbol_at_heading_size_is_rejected(self):
        assert ri._looks_like_heading("∑", 24, 10) is False
        assert ri._looks_like_heading("∏", 24, 10) is False

    def test_short_low_alpha_fragment_is_rejected(self):
        assert ri._looks_like_heading("Iδ", 24, 10) is False

    def test_real_headings_are_accepted(self):
        assert ri._looks_like_heading("Wyniki", 18, 10) is True
        assert ri._looks_like_heading("1. Introduction", 18, 10) is True

    def test_single_word_heading_still_accepted(self):
        # Abstract/Conclusion/References etc. are common single-word
        # headings — a naive word-count-only check must not reject these.
        assert ri._looks_like_heading("Abstract", 18, 10) is True


class TestParseDocx:
    def _docx_bytes(self) -> bytes:
        d = DocxDocument()
        d.add_heading("Rozdzial 1", level=1)
        d.add_paragraph(" ".join(f"abc{i}" for i in range(50)))
        d.add_heading("Rozdzial 2", level=1)
        d.add_paragraph(" ".join(f"xyz{i}" for i in range(50)))
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()

    def test_heading_styles_become_section_boundaries(self):
        paragraphs, page_count = ri.parse_docx(self._docx_bytes())
        assert page_count == 0
        assert {p.heading for p in paragraphs} == {"Rozdzial 1", "Rozdzial 2"}

    def test_body_paragraphs_are_not_treated_as_headings(self):
        paragraphs, _ = ri.parse_docx(self._docx_bytes())
        assert all(p.text not in ("Rozdzial 1", "Rozdzial 2") for p in paragraphs)


class TestParseText:
    def test_markdown_headings_split_sections(self):
        md = b"# Wstep\n\nAkapit pierwszy.\n\n# Wyniki\n\nAkapit drugi z wynikami."
        paragraphs, page_count = ri.parse_text(md)
        assert page_count == 0
        assert [p.heading for p in paragraphs] == ["Wstep", "Wyniki"]

    def test_plain_text_without_headings_has_none_section(self):
        paragraphs, _ = ri.parse_text(b"Just one paragraph.\n\nAnd another.")
        assert all(p.heading is None for p in paragraphs)


class TestChunkParagraphs:
    def test_never_crosses_a_section_boundary(self):
        pdf = _make_pdf([("Wprowadzenie", 200, "slowo"), ("Metodologia badan", 200, "inneslowo")])
        paragraphs, _ = ri.parse_pdf(pdf)
        chunks = ri.chunk_paragraphs(paragraphs, target_words=50, overlap_words=10)
        sections = {c.section for c in chunks}
        assert sections == {"Wprowadzenie", "Metodologia badan"}
        # every word in a "Wprowadzenie" chunk actually came from that section
        for c in chunks:
            if c.section == "Wprowadzenie":
                assert all(w.startswith("slowo") for w in c.text.split())

    def test_consecutive_chunks_overlap_by_the_configured_amount(self):
        pdf = _make_pdf([("Sekcja", 200, "w")])
        paragraphs, _ = ri.parse_pdf(pdf)
        chunks = ri.chunk_paragraphs(paragraphs, target_words=50, overlap_words=10)
        assert len(chunks) >= 2
        assert chunks[0].text.split()[-10:] == chunks[1].text.split()[:10]

    def test_empty_input_yields_no_chunks(self):
        assert ri.chunk_paragraphs([]) == []

    def test_ordinals_are_sequential(self):
        pdf = _make_pdf([("A", 100, "a"), ("B", 100, "b")])
        paragraphs, _ = ri.parse_pdf(pdf)
        chunks = ri.chunk_paragraphs(paragraphs, target_words=30, overlap_words=5)
        assert [c.ordinal for c in chunks] == list(range(len(chunks)))


class TestEmbeddingInput:
    def test_includes_filename_and_section(self):
        chunk = ri.Chunk(ordinal=0, page=1, section="Wyniki", text="tresc fragmentu", word_count=2)
        result = ri.embedding_input(chunk, "praca.pdf")
        assert result == "[praca.pdf > Wyniki] tresc fragmentu"

    def test_omits_section_when_none(self):
        chunk = ri.Chunk(ordinal=0, page=None, section=None, text="tresc", word_count=1)
        assert ri.embedding_input(chunk, "notes.txt") == "[notes.txt] tresc"


class TestParseDispatch:
    def test_unsupported_format_raises(self):
        with pytest.raises(ValueError):
            ri.parse(b"data", "exe")


class TestCleanText:
    def test_strips_nul_bytes(self):
        assert ri._clean_text("hello\x00world") == "helloworld"

    def test_strips_other_c0_control_chars(self):
        assert ri._clean_text("a\x01b\x1fc\x7fd") == "abcd"

    def test_preserves_whitespace(self):
        assert ri._clean_text("a\tb\nc\rd") == "a\tb\nc\rd"

    def test_pdf_parsing_never_yields_nul_bytes(self):
        # Postgres's text columns reject embedded NUL bytes outright — this
        # is a regression guard for that failure mode, not just a unit test
        # of _clean_text in isolation. (DOCX doesn't need the same guard:
        # XML itself disallows embedded NULs, so python-docx/lxml refuse to
        # even construct one — parse_docx still calls _clean_text out of
        # caution, but there's no legitimate DOCX to reproduce this with.)
        pdf = _make_pdf([("Nag\x00owek", 20, "sl\x00owo")])
        paragraphs, _ = ri.parse_pdf(pdf)
        assert all("\x00" not in p.text for p in paragraphs)
        assert all("\x00" not in (p.heading or "") for p in paragraphs)


class TestIngestDocument:
    async def test_happy_path_embeds_stores_and_marks_done(self, monkeypatch):
        import rag_embed
        import rag_store

        embed_mock = AsyncMock(return_value=[[0.1, 0.2]])
        insert_mock = AsyncMock()
        status_mock = AsyncMock()
        monkeypatch.setattr(rag_embed, "embed_texts", embed_mock)
        monkeypatch.setattr(rag_store, "insert_chunks", insert_mock)
        monkeypatch.setattr(rag_store, "set_document_status", status_mock)

        md = b"# Wstep\n\nJedno zdanie tresci."
        await ri.ingest_document("doc-1", "alice", "notes.md", "md", md)

        embed_mock.assert_awaited_once()
        insert_mock.assert_awaited_once()
        status_mock.assert_awaited_once_with("doc-1", "done", page_count=0)

    async def test_no_extractable_text_marks_error_without_embedding(self, monkeypatch):
        import rag_embed
        import rag_store

        embed_mock = AsyncMock()
        status_mock = AsyncMock()
        monkeypatch.setattr(rag_embed, "embed_texts", embed_mock)
        monkeypatch.setattr(rag_store, "set_document_status", status_mock)

        await ri.ingest_document("doc-1", "alice", "empty.txt", "txt", b"")

        embed_mock.assert_not_awaited()
        status_mock.assert_awaited_once_with("doc-1", "error", error="No extractable text found")

    async def test_embedding_failure_marks_error_and_reraises(self, monkeypatch):
        import rag_embed
        import rag_store

        monkeypatch.setattr(rag_embed, "embed_texts", AsyncMock(side_effect=RuntimeError("ollama down")))
        status_mock = AsyncMock()
        monkeypatch.setattr(rag_store, "set_document_status", status_mock)

        with pytest.raises(RuntimeError, match="ollama down"):
            await ri.ingest_document("doc-1", "alice", "notes.md", "md", b"# H\n\ntext")

        status_mock.assert_awaited_once_with("doc-1", "error", error="ollama down")


class TestOcrFallback:
    """OCR runs through PyMuPDF's built-in Tesseract integration, which
    needs the system tesseract-ocr binary — not assumed present in this test
    environment, so these mock fitz.Page.get_textpage_ocr() rather than
    exercising real OCR."""

    def test_has_text_true_only_with_non_blank_spans(self):
        assert ri._has_text([]) is False
        assert ri._has_text([{"lines": [{"spans": [{"text": "   "}]}]}]) is False
        assert ri._has_text([{"lines": [{"spans": [{"text": "hello"}]}]}]) is True

    def test_blank_page_falls_back_to_ocr(self, monkeypatch):
        calls = []

        def fake_get_textpage_ocr(self, language=None, dpi=None, full=None):
            calls.append((language, dpi, full))
            return "FAKE_TEXTPAGE"

        original_get_text = fitz.Page.get_text

        def fake_get_text(self, *args, **kwargs):
            if kwargs.get("textpage") == "FAKE_TEXTPAGE":
                return {"blocks": [{"lines": [{"spans": [{"text": "OCRed text", "size": 12.0}]}]}]}
            return original_get_text(self, *args, **kwargs)

        monkeypatch.setattr(fitz.Page, "get_textpage_ocr", fake_get_textpage_ocr)
        monkeypatch.setattr(fitz.Page, "get_text", fake_get_text)

        doc = fitz.open()
        doc.new_page()  # no inserted text — no text layer at all
        pdf_bytes = doc.tobytes()
        doc.close()

        paragraphs, page_count = ri.parse_pdf(pdf_bytes)
        assert page_count == 1
        assert any("OCRed text" in p.text for p in paragraphs)
        assert calls == [(ri.RAG_OCR_LANGUAGES, ri.RAG_OCR_DPI, True)]

    def test_ocr_failure_leaves_page_textless_without_crashing(self, monkeypatch):
        def raising_get_textpage_ocr(self, **kwargs):
            raise RuntimeError("tesseract not found")

        monkeypatch.setattr(fitz.Page, "get_textpage_ocr", raising_get_textpage_ocr)

        doc = fitz.open()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        paragraphs, page_count = ri.parse_pdf(pdf_bytes)
        assert page_count == 1
        assert paragraphs == []

    def test_page_with_real_text_never_triggers_ocr(self, monkeypatch):
        calls = []
        monkeypatch.setattr(fitz.Page, "get_textpage_ocr", lambda self, **kw: calls.append(1))

        pdf = _make_pdf([("Heading", 20, "word")])
        paragraphs, _ = ri.parse_pdf(pdf)
        assert paragraphs  # got real text
        assert calls == []
